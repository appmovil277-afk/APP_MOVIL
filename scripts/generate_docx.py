"""
Generador de documentación técnica Word para TallerFlow Muebles.
Produce un documento .docx profesional con mínimo 10 hojas de contenido
referente al código fuente y la arquitectura de la aplicación móvil.

Uso:
    python scripts/generate_docx.py
"""

from pathlib import Path
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Documentacion-Tecnica-TallerFlow-Muebles.docx"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def set_cell_shading(cell, color_hex: str):
    """Apply background shading to a table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a formatted table with header row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_shading(cell, "0F766E")
    # Data
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, "EAF6F6")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph("")
    return table


def add_code_block(doc, code: str, language: str = ""):
    """Add a code block styled paragraph."""
    if language:
        doc.add_paragraph(f"[{language}]").runs[0].font.color.rgb = RGBColor(100, 100, 100)
    p = doc.add_paragraph()
    run = p.add_run(code)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)


def section_break(doc):
    """Force a page break."""
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Document sections
# ---------------------------------------------------------------------------

def cover_page(doc):
    """Portada profesional."""
    for _ in range(6):
        doc.add_paragraph("")
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TallerFlow Muebles")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = RGBColor(15, 118, 110)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Documentación Técnica de la Aplicación Móvil")
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph("")
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = desc.add_run(
        "Sistema integral de gestión de producción de muebles a medida\n"
        "Expo · React Native · Firebase · TypeScript"
    )
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(120, 120, 120)

    for _ in range(4):
        doc.add_paragraph("")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Versión: 1.0.0\n"
        "Plataformas: iOS · Android · Web\n"
        "Lenguaje: TypeScript (strict)\n"
        "Framework: Expo SDK 55 / React Native 0.83"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)
    section_break(doc)


def table_of_contents(doc):
    """Tabla de contenidos manual."""
    doc.add_heading("Tabla de Contenidos", level=1)
    items = [
        "1. Resumen Ejecutivo",
        "2. Arquitectura del Sistema",
        "3. Estructura del Proyecto",
        "4. Stack Tecnológico y Dependencias",
        "5. Modelo de Datos (TypeScript Types)",
        "6. Componente Principal — App.tsx",
        "7. Servicios (services/)",
        "8. Utilidades (utils/)",
        "9. Datos Iniciales y Configuración",
        "10. Firebase: Reglas, Índices y Cloud Functions",
        "11. Sistema de Diseño y Temas",
        "12. Flujo de Trabajo de un Pedido",
        "13. Sistema de Roles y Permisos",
        "14. Sistema de Notificaciones",
        "15. Generación de PDF y Correos",
        "16. Persistencia y Sincronización",
        "17. Scripts y Herramientas de Desarrollo",
        "18. Catálogo de 250 Mejoras",
        "19. Guía de Despliegue",
        "20. Glosario Técnico",
    ]
    for item in items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
    section_break(doc)


def ch01_resumen(doc):
    doc.add_heading("1. Resumen Ejecutivo", level=1)
    doc.add_paragraph(
        "TallerFlow Muebles es una aplicación móvil y web multiplataforma desarrollada con "
        "Expo (SDK 55) y React Native 0.83 que digitaliza la operación completa de un taller "
        "de muebles a medida. La aplicación cubre el ciclo de vida de un pedido desde la "
        "captura comercial hasta la entrega al cliente final, pasando por diseño técnico, "
        "carpintería, tapicería, control de calidad y despacho."
    )
    doc.add_paragraph(
        "El sistema implementa un modelo de roles donde cada miembro del equipo solo accede "
        "a las etapas que le corresponden, mientras que el administrador tiene control total "
        "sobre la operación. Incluye persistencia dual (local + Firestore), generación de PDF, "
        "envío de correos, notificaciones push, y un log de auditoría automático."
    )
    doc.add_heading("Problema que resuelve", level=2)
    add_styled_table(doc,
        ["Problema", "Solución TallerFlow"],
        [
            ["Pedidos en papel o WhatsApp", "Formulario digital estructurado con referencia única"],
            ["Sin visibilidad del estado de producción", "Dashboard con KPIs y progreso en tiempo real"],
            ["Pagos desorganizados al equipo", "Seguimiento de pagos por etapa con reglas configurables"],
            ["Comunicación informal entre áreas", "Centro de notificaciones multicanal con historial"],
            ["Sin trazabilidad de acciones", "Log de auditoría automático en cada mutación"],
            ["Documentación manual para clientes", "Generación de PDF y envío por correo automático"],
        ])
    doc.add_heading("Datos clave del proyecto", level=2)
    add_styled_table(doc,
        ["Métrica", "Valor"],
        [
            ["Líneas de código (App.tsx)", "~3,500+"],
            ["Archivos TypeScript", "12"],
            ["Tipos definidos", "20+"],
            ["Roles del sistema", "7"],
            ["Vistas/pestañas", "10"],
            ["Mejoras catalogadas", "250"],
            ["Empleados demo", "7"],
            ["Pedidos demo", "3"],
        ])
    section_break(doc)


def ch02_arquitectura(doc):
    doc.add_heading("2. Arquitectura del Sistema", level=1)
    doc.add_heading("Diagrama de capas", level=2)
    add_code_block(doc, """
┌─────────────────────────────────────────────────────────┐
│              CAPA DE PRESENTACIÓN                        │
│  App.tsx (monolito React con Context + Estado local)     │
│  10 pestañas: Dashboard, Solicitudes, Flujo, Historial, │
│  Notificaciones, Clientes/PDF, Equipo, Mejoras,         │
│  Centro Avanzado, Ajustes                                │
├─────────────────────────────────────────────────────────┤
│              CAPA DE LÓGICA DE NEGOCIO                   │
│  utils/order.ts    → Creación de pedidos, flujo          │
│  utils/format.ts   → Moneda COP, fechas, etiquetas      │
│  utils/backend.ts  → IDs, audit, sync queue, salud      │
├─────────────────────────────────────────────────────────┤
│              CAPA DE SERVICIOS                           │
│  services/repository.ts   → Persistencia dual            │
│  services/firebase.ts     → Firestore CRUD               │
│  services/notifications.ts→ Push + web notifications     │
│  services/pdf.ts          → HTML→PDF + compartir         │
├─────────────────────────────────────────────────────────┤
│              CAPA DE DATOS                               │
│  AsyncStorage (local)  ←→  Firestore (nube)              │
│  data/seed.ts (datos iniciales)                          │
├─────────────────────────────────────────────────────────┤
│              CAPA BACKEND SERVERLESS                     │
│  Firebase Cloud Functions (Node.js 20)                   │
│  Trigger: mailQueue → Nodemailer → Gmail SMTP            │
└─────────────────────────────────────────────────────────┘
    """.strip(), "Texto plano")

    doc.add_heading("Patrones de diseño implementados", level=2)
    add_styled_table(doc,
        ["Patrón", "Implementación", "Ubicación"],
        [
            ["Context + Local State", "Estado global workspace + ThemeContext", "App.tsx"],
            ["Optimistic Updates", "Cambios locales inmediatos, sync en background", "App.tsx → mutate()"],
            ["Factory Pattern", "Funciones creadoras de entidades tipadas", "utils/backend.ts, utils/order.ts"],
            ["Observer", "Auditoría y notificaciones en cada mutación", "App.tsx → mutate()"],
            ["Repository", "Abstracción de almacenamiento dual", "services/repository.ts"],
            ["Role-Based Access", "Vistas filtradas por activeRole", "App.tsx"],
        ])

    doc.add_heading("Flujo de mutación de estado", level=2)
    doc.add_paragraph(
        "Todas las mutaciones del estado pasan por la función mutate() que orquesta:"
    )
    steps = [
        "1. Actualiza el objeto workspace con la transformación proporcionada.",
        "2. Crea una entrada en el log de auditoría (AuditEntry).",
        "3. Añade un trabajo a la cola de sincronización (SyncJob).",
        "4. Encola un correo electrónico si se proporciona mailItem.",
        "5. Emite una notificación local (AppNotification).",
        "6. Persiste en AsyncStorage con debounce de 300ms.",
        "7. Sincroniza con Firestore en background si está configurado.",
    ]
    for s in steps:
        doc.add_paragraph(s, style="List Number")
    section_break(doc)


def ch03_estructura(doc):
    doc.add_heading("3. Estructura del Proyecto", level=1)
    doc.add_paragraph(
        "El proyecto sigue una estructura modular organizada por responsabilidad. "
        "El componente principal App.tsx contiene toda la interfaz y la lógica de estado, "
        "mientras que los módulos en src/ proveen servicios, utilidades y constantes."
    )
    add_code_block(doc, """
app-movil/
├── App.tsx                    # Componente principal (UI + lógica + estado)
├── index.ts                   # Entry point Expo (registerRootComponent)
├── app.json                   # Configuración Expo
├── package.json               # Dependencias y scripts npm
├── tsconfig.json              # TypeScript strict mode
├── firebase.json              # Configuración Firebase CLI
├── firestore.rules            # Reglas de seguridad Firestore
├── firestore.indexes.json     # Índices compuestos
├── .env.example               # Plantilla de variables de entorno
│
├── src/
│   ├── types.ts               # Definiciones TypeScript del sistema
│   ├── constants/
│   │   ├── theme.ts           # Temas oscuro/claro, colores por rol
│   │   └── improvements.ts    # 250 mejoras categorizadas
│   ├── data/
│   │   └── seed.ts            # Datos iniciales: empleados y pedidos
│   ├── services/
│   │   ├── firebase.ts        # Conexión y CRUD Firestore
│   │   ├── notifications.ts   # Push local + Web Notifications
│   │   ├── pdf.ts             # HTML→PDF + compartir + email
│   │   └── repository.ts      # Persistencia dual local/cloud
│   └── utils/
│       ├── backend.ts         # uid(), factories, salud del sistema
│       ├── format.ts          # Moneda COP, fechas, etiquetas
│       └── order.ts           # Blueprints, builder, progreso
│
├── firebase/functions/
│   ├── index.js               # Cloud Function: auto-email
│   └── package.json           # firebase-admin + nodemailer
│
├── scripts/
│   ├── generate_docx.py       # Este script, genera este documento
│   └── serve-web.ps1          # Servidor para build web exportado
│
├── docs/
│   ├── documentacion.md       # Documentación funcional
│   └── firebase-setup.md      # Guía de configuración Firebase
│
└── assets/                    # Iconos y splash screens
    """.strip(), "Texto plano")

    doc.add_heading("Descripción de cada archivo fuente", level=2)
    add_styled_table(doc,
        ["Archivo", "Líneas aprox.", "Responsabilidad"],
        [
            ["App.tsx", "3,500+", "Interfaz completa, estado global, navegación, lógica de negocio"],
            ["src/types.ts", "120+", "Todas las interfaces y tipos del dominio"],
            ["src/constants/theme.ts", "200+", "Paletas de colores, gradientes, sombras, tipografía"],
            ["src/constants/improvements.ts", "500+", "250 mejoras en 10 categorías"],
            ["src/data/seed.ts", "300+", "7 empleados, 3 pedidos, notificaciones demo"],
            ["src/services/firebase.ts", "120+", "Inicialización Firebase, load/save workspace"],
            ["src/services/notifications.ts", "60+", "Permisos, handler, envío push/web"],
            ["src/services/pdf.ts", "150+", "Template HTML, generación PDF, share, email"],
            ["src/services/repository.ts", "100+", "AsyncStorage + Firestore, sesión, tema"],
            ["src/utils/backend.ts", "80+", "Generadores de ID, factories, salud"],
            ["src/utils/format.ts", "70+", "formatCurrency, formatDate, etiquetas"],
            ["src/utils/order.ts", "200+", "Blueprints de etapas, builder de pedidos"],
            ["firebase/functions/index.js", "100+", "Cloud Function mailQueueCreated"],
        ])
    section_break(doc)


def ch04_stack(doc):
    doc.add_heading("4. Stack Tecnológico y Dependencias", level=1)
    doc.add_heading("Dependencias de producción", level=2)
    add_styled_table(doc,
        ["Paquete", "Versión", "Propósito"],
        [
            ["expo", "55.0.5", "Framework React Native multiplataforma"],
            ["react", "19.2.0", "Biblioteca de interfaz de usuario"],
            ["react-native", "0.83.2", "Renderizado nativo iOS y Android"],
            ["react-native-web", "0.21.0", "Soporte para plataforma web"],
            ["react-dom", "19.2.0", "Renderizado DOM para web"],
            ["expo-linear-gradient", "15.0.2", "Fondos con gradientes"],
            ["expo-blur", "14.1.5", "Efecto glassmorphism (desenfoque)"],
            ["@expo/vector-icons", "15.0.2", "Todos los iconos (Ionicons)"],
            ["@react-native-async-storage/async-storage", "2.2.0", "Almacenamiento local key-value"],
            ["@react-native-picker/picker", "2.11.4", "Componente dropdown nativo"],
            ["expo-print", "~13.2.3", "Generación de PDF desde HTML"],
            ["expo-sharing", "~13.0.1", "Hoja de compartir nativa"],
            ["expo-mail-composer", "~14.0.2", "Composición de correo con adjuntos"],
            ["expo-notifications", "~0.32.2", "Notificaciones push locales"],
            ["firebase", "12.10.0", "SDK web para Firestore"],
        ])
    doc.add_heading("Dependencias de desarrollo", level=2)
    add_styled_table(doc,
        ["Paquete", "Versión", "Propósito"],
        [
            ["typescript", "5.9.2", "Tipado estático (modo estricto)"],
            ["@types/react", "19.2.2", "Tipos para React"],
            ["@babel/core", "7.25.2", "Transpilador JavaScript"],
        ])
    doc.add_heading("Dependencias Cloud Functions", level=2)
    add_styled_table(doc,
        ["Paquete", "Versión", "Propósito"],
        [
            ["firebase-admin", "13.5.0", "SDK Admin para Firestore server-side"],
            ["firebase-functions", "6.5.0", "Framework de Cloud Functions v2"],
            ["nodemailer", "6.9.16", "Envío de correos vía SMTP"],
        ])
    section_break(doc)


def ch05_modelo_datos(doc):
    doc.add_heading("5. Modelo de Datos (TypeScript Types)", level=1)
    doc.add_paragraph(
        "Todas las entidades del sistema están definidas en src/types.ts con TypeScript estricto. "
        "A continuación se documenta cada tipo con sus campos, tipos y descripción."
    )

    doc.add_heading("5.1 Tipos de rol y estado", level=2)
    add_styled_table(doc,
        ["Tipo", "Valores posibles", "Descripción"],
        [
            ["RoleId", "comercial, administracion, diseno, carpinteria, tapiceria, calidad, despacho", "Identificador del rol del empleado"],
            ["OrderStatus", "pending_approval, approved, in_progress, blocked, completed", "Estado del ciclo de vida del pedido"],
            ["StageStatus", "pending, active, completed, blocked", "Estado de una etapa individual"],
            ["Priority", "Alta, Media, Baja", "Prioridad del pedido"],
            ["PayoutRule", "Al registrar, Al aprobar, Al iniciar, Al finalizar, Contra entrega", "Momento en que se paga la etapa"],
            ["NotificationChannel", "Push, Email, In-app, WhatsApp", "Canal de envío de notificación"],
        ])

    doc.add_heading("5.2 Client (cliente)", level=2)
    add_styled_table(doc,
        ["Campo", "Tipo", "Descripción"],
        [
            ["name", "string", "Nombre completo del cliente"],
            ["email", "string", "Correo electrónico"],
            ["phone", "string", "Teléfono de contacto"],
            ["city", "string", "Ciudad de ubicación"],
        ])

    doc.add_heading("5.3 Employee (empleado)", level=2)
    add_styled_table(doc,
        ["Campo", "Tipo", "Descripción"],
        [
            ["id", "string", "Identificador único (emp-XXXXXXXX)"],
            ["name", "string", "Nombre completo"],
            ["role", "RoleId", "Rol principal asignado"],
            ["email", "string", "Correo electrónico laboral"],
            ["phone", "string", "Teléfono"],
            ["avatarUri", "string?", "URL de foto de perfil (opcional)"],
            ["accent", "string", "Color de acento para UI (#hex)"],
            ["baseRate", "number", "Tarifa base por etapa (COP)"],
            ["availability", "'Disponible' | 'Ocupado' | 'Ausente'", "Estado de disponibilidad"],
            ["specialties", "string[]", "Lista de especialidades"],
            ["performance", "number", "Rendimiento 0-100 (porcentaje)"],
            ["paymentMethods", "string[]?", "Métodos: Nequi, Daviplata, Bancolombia"],
            ["paymentAccounts", "Record<string, string>?", "Cuentas por método de pago"],
        ])

    doc.add_heading("5.4 WorkflowStage (etapa)", level=2)
    add_styled_table(doc,
        ["Campo", "Tipo", "Descripción"],
        [
            ["id", "string", "Identificador único (stg-XXXXXXXX)"],
            ["title", "string", "Nombre de la etapa"],
            ["role", "RoleId", "Rol responsable"],
            ["assigneeId", "string", "ID del empleado asignado"],
            ["status", "StageStatus", "Estado actual de la etapa"],
            ["payout", "number", "Monto de pago (COP)"],
            ["payoutRule", "PayoutRule", "Regla de momento de pago"],
            ["note", "string", "Nota adicional del administrador"],
            ["startedAt", "string?", "Timestamp ISO de inicio"],
            ["completedAt", "string?", "Timestamp ISO de finalización"],
            ["paidAt", "string?", "Timestamp ISO de pago"],
        ])

    doc.add_heading("5.5 FurnitureOrder (pedido)", level=2)
    add_styled_table(doc,
        ["Campo", "Tipo", "Descripción"],
        [
            ["id", "string", "Identificador único (ord-XXXXXXXX)"],
            ["reference", "string", "Referencia secuencial (#1, #2...)"],
            ["createdAt", "string", "Timestamp ISO de creación"],
            ["updatedAt", "string", "Timestamp ISO de última actualización"],
            ["createdByEmployeeId", "string", "ID del empleado creador"],
            ["client", "Client", "Datos del cliente"],
            ["furnitureName", "string", "Nombre del mueble"],
            ["material", "string", "Material seleccionado"],
            ["size", "string", "Dimensiones del mueble"],
            ["needsCushions", "boolean", "Si requiere almohadas/cojines"],
            ["estimatedCost", "number", "Costo estimado por comercial (COP)"],
            ["dueDate", "string", "Fecha límite de entrega"],
            ["notes", "string", "Notas adicionales"],
            ["status", "OrderStatus", "Estado del ciclo de vida"],
            ["priority", "Priority", "Prioridad calculada"],
            ["pricing", "PricingInfo", "Costos: producción, final, margen"],
            ["stages", "WorkflowStage[]", "Etapas del flujo de trabajo"],
            ["events", "OrderEvent[]", "Historial de eventos"],
            ["tags", "string[]", "Etiquetas: material, almohadas, tamaño"],
        ])

    doc.add_heading("5.6 AppNotification", level=2)
    add_styled_table(doc,
        ["Campo", "Tipo", "Descripción"],
        [
            ["id", "string", "Identificador único (not-XXXXXXXX)"],
            ["title", "string", "Título de la notificación"],
            ["body", "string", "Cuerpo descriptivo"],
            ["createdAt", "string", "Timestamp ISO"],
            ["priority", "Priority", "Nivel de prioridad"],
            ["channels", "NotificationChannel[]", "Canales de envío"],
            ["orderId", "string?", "Pedido relacionado"],
            ["recipients", "string[]", "Nombres de destinatarios"],
            ["read", "boolean", "Si fue leída"],
            ["actionLabel", "string?", "Etiqueta de acción sugerida"],
        ])

    doc.add_heading("5.7 Entidades operativas", level=2)
    doc.add_paragraph("SyncJob — Cola de sincronización:")
    add_styled_table(doc,
        ["Campo", "Tipo", "Descripción"],
        [
            ["id", "string", "sync-XXXXXXXX"],
            ["entityId", "string", "ID de la entidad afectada"],
            ["entityType", "string", "workspace | order | notification | mailQueue"],
            ["action", "string", "Acción realizada (approve_order, create_order, etc.)"],
            ["status", "string", "pending | synced | failed"],
            ["createdAt", "string", "Timestamp ISO"],
            ["message", "string", "Descripción legible de la operación"],
        ])

    doc.add_paragraph("MailQueueItem — Cola de correos:")
    add_styled_table(doc,
        ["Campo", "Tipo", "Descripción"],
        [
            ["id", "string", "mail-XXXXXXXX"],
            ["orderId", "string", "Pedido relacionado"],
            ["recipient", "string", "Email del destinatario"],
            ["subject", "string", "Asunto del correo"],
            ["body", "string", "Cuerpo del correo"],
            ["status", "string", "pending | prepared | sent | failed"],
            ["pdfRequested", "boolean", "Si se pidió adjuntar PDF"],
        ])

    doc.add_paragraph("AuditEntry — Log de auditoría:")
    add_styled_table(doc,
        ["Campo", "Tipo", "Descripción"],
        [
            ["id", "string", "audit-XXXXXXXX"],
            ["action", "string", "Acción: create_order, approve_order, etc."],
            ["actorId", "string", "ID del empleado que realizó la acción"],
            ["detail", "string", "Descripción detallada"],
            ["orderId", "string?", "Pedido relacionado"],
            ["createdAt", "string", "Timestamp ISO"],
        ])

    doc.add_heading("5.8 WorkspaceState (estado raíz)", level=2)
    add_styled_table(doc,
        ["Campo", "Tipo", "Descripción"],
        [
            ["employees", "Employee[]", "Lista de todos los empleados"],
            ["orders", "FurnitureOrder[]", "Lista de todos los pedidos"],
            ["notifications", "AppNotification[]", "Todas las notificaciones"],
            ["syncQueue", "SyncJob[]", "Cola de sincronización"],
            ["mailQueue", "MailQueueItem[]", "Cola de correos"],
            ["auditLog", "AuditEntry[]", "Registro de auditoría"],
            ["notificationSettings", "NotificationSettings", "Configuración de notificaciones"],
            ["lastSyncedAt", "string?", "Último sync exitoso con Firebase"],
        ])
    section_break(doc)


def ch06_app_tsx(doc):
    doc.add_heading("6. Componente Principal — App.tsx", level=1)
    doc.add_paragraph(
        "App.tsx es el archivo central de la aplicación con aproximadamente 3,500 líneas de código. "
        "Contiene toda la interfaz de usuario, la lógica de estado y la navegación. "
        "Utiliza React Context para el tema y estado local para los datos del workspace."
    )

    doc.add_heading("6.1 Contexto de tema (ThemeContext)", level=2)
    doc.add_paragraph(
        "Provee colores y estilos computados a toda la aplicación. Permite alternar entre "
        "tema oscuro (aurora boreal) y claro. Se persiste en AsyncStorage."
    )

    doc.add_heading("6.2 Flujo de autenticación", level=2)
    doc.add_paragraph(
        "Login con perfiles demo predefinidos. Al seleccionar un perfil, se almacena la sesión "
        "en AsyncStorage y se carga el workspace completo."
    )

    doc.add_heading("6.3 Pestañas de la aplicación", level=2)
    add_styled_table(doc,
        ["Pestaña", "ID interno", "Descripción", "Roles con acceso"],
        [
            ["Dashboard", "home", "KPIs, acciones rápidas, pedidos recientes, alertas", "Todos"],
            ["Solicitudes", "solicitudes", "Formulario de captura comercial con plantillas", "Comercial, Admin"],
            ["Administración", "administracion", "Redirección a Flujo (consolidado)", "Admin"],
            ["Flujo/Tareas", "flujo", "Pedidos activos, aprobar, iniciar/completar etapas", "Todos"],
            ["Historial", "historial", "Pedidos completados con indicadores de pago", "Todos"],
            ["Notificaciones", "notificaciones", "Centro de alertas con búsqueda y filtrado", "Todos"],
            ["Clientes/PDF", "clientes", "Generación de PDF, compartir, enviar email", "Todos"],
            ["Equipo", "equipo", "Directorio de empleados con métricas", "Todos"],
            ["Mejoras", "mejoras", "Catálogo de 250 mejoras categorizadas", "Todos"],
            ["Centro Avanzado", "mas", "Sync health, colas, auditoría, cambiar perfil", "Todos"],
            ["Ajustes", "ajustes", "Perfil, tema, métodos de pago, logout", "Todos"],
        ])

    doc.add_heading("6.4 Componentes UI internos", level=2)
    add_styled_table(doc,
        ["Componente", "Props principales", "Descripción"],
        [
            ["Stat", "label, value", "Tarjeta KPI con etiqueta y valor grande"],
            ["Chip", "label, active, onPress", "Badge toggleable con gradiente"],
            ["Panel", "title, subtitle, children", "Contenedor con efecto vidrio"],
            ["Field", "label, children", "Wrapper de input con etiqueta"],
            ["Tag", "label, color", "Etiqueta inline pequeña"],
            ["Quick", "label, icon, onPress", "Botón ghost compacto"],
            ["AvatarNode", "name, accent, isCurrent", "Círculo con iniciales + pulso"],
            ["StyledPicker", "selectedValue, onValueChange, items", "Dropdown estilizado teal"],
            ["SettingCard", "title, description, value, onToggle", "Toggle con descripción"],
            ["BottomTabBar", "tabs, activeTab, onPress, unread", "Barra inferior estilo iPhone"],
        ])

    doc.add_heading("6.5 Función mutate()", level=2)
    doc.add_paragraph(
        "Función central que orquesta todas las mutaciones de estado. Recibe un updater function, "
        "un mensaje descriptivo, una notificación opcional y metadata de backend. Automáticamente "
        "crea entradas de auditoría, trabajos de sync, correos encolados y notificaciones locales."
    )
    add_code_block(doc, """function mutate(
  updater: (current: WorkspaceState) => WorkspaceState,
  message?: string,
  notification?: AppNotification,
  backendMeta?: {
    action: string,
    detail: string,
    entityId: string,
    entityType: string,
    orderId?: string,
    mailItem?: MailQueueItem
  }
)""", "TypeScript")
    section_break(doc)


def ch07_servicios(doc):
    doc.add_heading("7. Servicios (services/)", level=1)

    doc.add_heading("7.1 firebase.ts — Conexión Firestore", level=2)
    doc.add_paragraph(
        "Gestiona la conexión con Firebase Firestore. Lee las credenciales de variables de entorno "
        "EXPO_PUBLIC_FIREBASE_*. Si alguna falta, opera en modo local sin errores."
    )
    doc.add_heading("Funciones principales:", level=3)
    add_styled_table(doc,
        ["Función", "Retorno", "Descripción"],
        [
            ["isFirebaseConfigured()", "boolean", "Verifica si todas las variables están presentes"],
            ["loadWorkspaceFromFirebase()", "WorkspaceState | null", "Carga todas las colecciones en paralelo"],
            ["saveWorkspaceToFirebase(state)", "boolean", "Escribe con batch atómico"],
            ["sanitizeForFirestore(data)", "any", "Limpia datos via JSON parse/stringify"],
        ])

    doc.add_heading("Colecciones Firestore:", level=3)
    add_styled_table(doc,
        ["Colección", "Documento", "Entidad"],
        [
            ["employees/{id}", "Auto-ID", "Employee"],
            ["orders/{id}", "Auto-ID", "FurnitureOrder"],
            ["notifications/{id}", "Auto-ID", "AppNotification"],
            ["syncQueue/{id}", "Auto-ID", "SyncJob"],
            ["mailQueue/{id}", "Auto-ID", "MailQueueItem"],
            ["auditLog/{id}", "Auto-ID", "AuditEntry"],
            ["settings/workspace", "Singleton", "NotificationSettings + lastSyncedAt"],
        ])

    doc.add_heading("7.2 repository.ts — Persistencia dual", level=2)
    doc.add_paragraph(
        "Implementa el patrón Repository con dos backends: AsyncStorage (siempre) y Firestore "
        "(si configurado). La lectura prioriza datos locales; la escritura va a ambos."
    )
    add_styled_table(doc,
        ["Función", "Descripción"],
        [
            ["loadWorkspace()", "Lee local → si vacío, intenta Firebase → si vacío, usa seed"],
            ["persistWorkspace(ws)", "Escribe a AsyncStorage + Firestore, marca sync jobs"],
            ["loadSession() / persistSession() / clearSession()", "Gestión de sesión activa"],
            ["loadThemeMode() / persistThemeMode()", "Preferencia de tema dark/light"],
        ])

    doc.add_heading("7.3 notifications.ts — Sistema de push", level=2)
    doc.add_paragraph(
        "Maneja permisos y envío de notificaciones tanto en móvil (Expo Notifications) como "
        "en web (Web Notifications API). Configurado con sonido, banner y badge activados."
    )

    doc.add_heading("7.4 pdf.ts — Generación de documentos", level=2)
    doc.add_paragraph(
        "Genera un documento HTML profesional con los datos del pedido y lo convierte a PDF. "
        "En web usa diálogo de impresión; en móvil genera archivo temporal."
    )
    add_styled_table(doc,
        ["Función", "Plataforma", "Descripción"],
        [
            ["buildOrderHtml(order, employees)", "Todas", "Genera HTML con hero, specs, costos, flujo"],
            ["createOrderPdf(order, employees)", "Móvil/Web", "HTML → PDF (archivo o diálogo)"],
            ["sharePdf(uri)", "Móvil", "Abre hoja de compartir nativa"],
            ["composeOrderEmail(order, uri)", "Todas", "Compone email con PDF adjunto"],
        ])
    section_break(doc)


def ch08_utilidades(doc):
    doc.add_heading("8. Utilidades (utils/)", level=1)

    doc.add_heading("8.1 backend.ts — Operaciones de backend", level=2)
    add_styled_table(doc,
        ["Función", "Descripción"],
        [
            ["uid(prefix)", "Genera ID único: prefix-XXXXXXXX (base36)"],
            ["createSyncJob(params)", "Crea trabajo de sincronización con estado 'pending'"],
            ["createAuditEntry(params)", "Crea entrada de auditoría con timestamp"],
            ["createMailQueueItem(params)", "Crea ítem de cola de correo con estado 'prepared'"],
            ["computeSystemHealth(ws)", "Calcula salud: healthy / warning / critical"],
        ])
    doc.add_heading("Criterios de salud del sistema:", level=3)
    add_styled_table(doc,
        ["Estado", "Condición"],
        [
            ["Critical", "Cola sync > 8 ó pedidos vencidos > 2"],
            ["Warning", "Cola sync > 3 ó pedidos vencidos > 0 ó correos pendientes > 3"],
            ["Healthy", "Todo dentro de rangos normales"],
        ])

    doc.add_heading("8.2 format.ts — Formateo y etiquetas", level=2)
    add_styled_table(doc,
        ["Función", "Ejemplo entrada", "Ejemplo salida"],
        [
            ["formatCurrency(3200000)", "3200000", "₡3,200,000"],
            ["formatDate('2026-03-14T15:45:00')", "ISO string", "14/03/2026, 03:45 p.m."],
            ["formatDateOnly('2026-03-14')", "ISO string", "14 de marzo de 2026"],
            ["getDaysRemaining('2026-03-16')", "fecha futura", "Faltan 2 día(s)"],
            ["getDaysRemaining('2026-03-12')", "fecha pasada", "Vencido hace 2 día(s)"],
            ["getOrderStatusLabel('in_progress')", "OrderStatus", "En producción"],
            ["getRoleLabel('carpinteria')", "RoleId", "Carpinteria"],
            ["clampNumber('abc', 0)", "texto inválido", "0"],
        ])

    doc.add_heading("8.3 order.ts — Lógica de pedidos", level=2)
    doc.add_paragraph("Blueprints de etapas predefinidas:")
    add_styled_table(doc,
        ["Etapa", "Rol", "Pago base (COP)", "Regla de pago", "Bloqueada"],
        [
            ["Solicitud comercial", "Comercial", "30,000", "Al registrar", "Sí"],
            ["Validación administrativa", "Administración", "45,000", "Al aprobar", "Sí"],
            ["Diseño técnico", "Diseño", "60,000", "Al finalizar", "No"],
            ["Carpintería", "Carpintería", "120,000", "Al finalizar", "No"],
            ["Tapicería", "Tapicería", "95,000", "Al finalizar", "No"],
            ["Control de calidad", "Calidad", "50,000", "Al finalizar", "No"],
            ["Despacho", "Despacho", "40,000", "Contra entrega", "No"],
        ])

    doc.add_heading("Funciones de flujo:", level=3)
    add_styled_table(doc,
        ["Función", "Descripción"],
        [
            ["buildOrderFromDraft()", "Construye FurnitureOrder desde el formulario comercial"],
            ["getCurrentExecutableStage()", "Primera etapa pendiente donde todas las anteriores están completadas"],
            ["getActiveStage()", "Etapa activa actual o siguiente ejecutable"],
            ["deriveOrderStatus()", "Calcula estado del pedido desde sus etapas"],
            ["getOrderProgress()", "Porcentaje completado (etapas completadas / total × 100)"],
            ["createOrderEvent()", "Crea evento de historial con timestamp"],
            ["createNotification()", "Crea notificación In-app + Push + Email"],
            ["updateMargins()", "Recalcula margen = finalPrice - productionCost"],
        ])
    section_break(doc)


def ch09_datos_config(doc):
    doc.add_heading("9. Datos Iniciales y Configuración", level=1)

    doc.add_heading("9.1 Empleados de demostración (seed.ts)", level=2)
    add_styled_table(doc,
        ["Nombre", "Rol", "Tarifa base", "Rendimiento", "Disponibilidad", "Especialidades"],
        [
            ["Laura Díaz", "Comercial", "$30,000", "92%", "Disponible", "Ventas, Atención al cliente"],
            ["Mateo Ruiz", "Administración", "$45,000", "89%", "Disponible", "Finanzas, Gestión"],
            ["Camila Soto", "Diseño", "$60,000", "94%", "Disponible", "CAD, Diseño interior"],
            ["Javier Peña", "Carpintería", "$120,000", "87%", "Ocupado", "Madera maciza, Ensamblaje"],
            ["Daniela Ríos", "Tapicería", "$95,000", "91%", "Disponible", "Telas, Acabados"],
            ["Sofía Lara", "Calidad", "$50,000", "96%", "Disponible", "Inspección, Normas"],
            ["Esteban Cruz", "Despacho", "$40,000", "88%", "Disponible", "Logística, Embalaje"],
        ])

    doc.add_heading("9.2 Pedidos de demostración", level=2)
    add_styled_table(doc,
        ["Referencia", "Mueble", "Cliente", "Ciudad", "Precio final", "Estado"],
        [
            ["#1", "Sofá modular Oslo", "Ana Mejía", "Bogotá", "COP 4,680,000", "En producción (diseño activo)"],
            ["#2", "Poltrona Verona", "Carlos Varela", "Medellín", "COP 2,240,000", "Pendiente de aprobación"],
            ["#3", "Cabecero Niza", "Mariana Torres", "Cali", "COP 3,210,000", "En producción (tapicería)"],
        ])

    doc.add_heading("9.3 Configuración Expo (app.json)", level=2)
    add_styled_table(doc,
        ["Parámetro", "Valor"],
        [
            ["Nombre de la app", "TallerFlow Muebles"],
            ["Slug", "tallerflow-muebles"],
            ["Versión", "1.0.0"],
            ["Orientación", "Portrait (vertical)"],
            ["Esquema URL", "tallerflow://"],
            ["Soporte tablet iOS", "Sí"],
            ["Iconos adaptativos Android", "Sí (foreground + background + monochrome)"],
            ["Plugins nativos", "expo-sharing, expo-mail-composer, expo-notifications"],
        ])

    doc.add_heading("9.4 TypeScript (tsconfig.json)", level=2)
    doc.add_paragraph(
        "Extiende la configuración base de Expo con modo estricto habilitado. Esto garantiza "
        "máxima seguridad de tipos: sin any implícito, sin null no verificado, retornos "
        "obligatorios y verificación de switch exhaustiva."
    )
    section_break(doc)


def ch10_firebase(doc):
    doc.add_heading("10. Firebase: Reglas, Índices y Cloud Functions", level=1)

    doc.add_heading("10.1 Reglas de seguridad Firestore", level=2)
    doc.add_paragraph(
        "Las reglas actuales permiten lectura y escritura a cualquier usuario autenticado "
        "(request.auth != null). Están definidas en firestore.rules."
    )
    add_code_block(doc, """rules_version = '2';
service cloud.firestore {
  match /databases/{database}/documents {
    function signedIn() {
      return request.auth != null;
    }
    match /employees/{employeeId} { allow read, write: if signedIn(); }
    match /orders/{orderId} { allow read, write: if signedIn(); }
    match /notifications/{notificationId} { allow read, write: if signedIn(); }
    match /syncQueue/{syncId} { allow read, write: if signedIn(); }
    match /mailQueue/{mailId} { allow read, write: if signedIn(); }
    match /auditLog/{auditId} { allow read, write: if signedIn(); }
    match /settings/{settingId} { allow read, write: if signedIn(); }
  }
}""", "JavaScript")

    doc.add_heading("10.2 Índices compuestos", level=2)
    add_styled_table(doc,
        ["Colección", "Campo 1", "Campo 2", "Propósito"],
        [
            ["orders", "status (ASC)", "dueDate (ASC)", "Pedidos por estado y fecha límite"],
            ["notifications", "read (ASC)", "createdAt (DESC)", "No leídas más recientes"],
            ["syncQueue", "status (ASC)", "createdAt (DESC)", "Trabajos pendientes recientes"],
            ["mailQueue", "status (ASC)", "createdAt (DESC)", "Correos pendientes recientes"],
        ])

    doc.add_heading("10.3 Cloud Function: mailQueueCreated", level=2)
    doc.add_paragraph(
        "Función serverless que se ejecuta automáticamente cuando se crea un documento en "
        "la colección mailQueue. Usa Nodemailer con SMTP de Gmail para enviar el correo."
    )
    doc.add_paragraph("Flujo de ejecución:")
    for step in [
        "1. Se detecta nuevo documento en mailQueue/{mailId}.",
        "2. Se extraen: recipient, subject, body, pdfRequested.",
        "3. Se crea un transporter Nodemailer con credenciales SMTP.",
        "4. Se genera el cuerpo HTML con diseño TallerFlow (gradiente teal).",
        "5. Se envía el correo al destinatario.",
        "6. Se actualiza el documento: status → 'sent' o 'failed'.",
        "7. Se registra el resultado en colección mailQueueLogs.",
    ]:
        doc.add_paragraph(step, style="List Number")
    section_break(doc)


def ch11_tema(doc):
    doc.add_heading("11. Sistema de Diseño y Temas", level=1)

    doc.add_heading("11.1 Tema oscuro (por defecto)", level=2)
    doc.add_paragraph(
        "Inspirado en aurora boreal con efecto glassmorphism. Fondos navy con gradientes "
        "de 3 capas, acentos de glow en azul, rosa, mint, índigo y ámbar."
    )
    add_styled_table(doc,
        ["Propiedad", "Valor", "Uso"],
        [
            ["primary", "#5ce1ff", "Color principal, botones, acentos"],
            ["bg (top)", "#030712", "Fondo superior del gradiente"],
            ["bg (bottom)", "#0a1528", "Fondo inferior del gradiente"],
            ["text", "#f3f7ff", "Texto principal claro"],
            ["glass overlay", "rgba(255,255,255,0.06)", "Paneles transparentes"],
            ["glow blue", "#5ce1ff22", "Aura luminosa azul"],
            ["glow pink", "#ff7aa522", "Aura luminosa rosa"],
            ["glow mint", "#7bffb622", "Aura luminosa verde"],
        ])

    doc.add_heading("11.2 Tema claro", level=2)
    add_styled_table(doc,
        ["Propiedad", "Valor", "Uso"],
        [
            ["primary", "#127ea2", "Color principal para tema claro"],
            ["bg", "#f0f7fb", "Fondo de la aplicación"],
            ["text", "#102038", "Texto principal oscuro"],
        ])

    doc.add_heading("11.3 Colores por rol", level=2)
    add_styled_table(doc,
        ["Rol", "Color hex", "Nombre del color"],
        [
            ["Comercial", "#c27d22", "Ámbar"],
            ["Administración", "#0f766e", "Teal"],
            ["Diseño", "#184e77", "Azul"],
            ["Carpintería", "#7f5539", "Marrón"],
            ["Tapicería", "#8d5a97", "Púrpura"],
            ["Calidad", "#577590", "Slate"],
            ["Despacho", "#3d405b", "Dark slate"],
        ])

    doc.add_heading("11.4 Colores por estado", level=2)
    add_styled_table(doc,
        ["Estado", "Color hex", "Nombre"],
        [
            ["Pendiente", "#8193b2", "Gris azulado"],
            ["Activo", "#5ce1ff", "Cian"],
            ["Completado", "#7bffb6", "Mint"],
            ["Bloqueado", "#ff7aa5", "Rosa"],
        ])
    section_break(doc)


def ch12_flujo(doc):
    doc.add_heading("12. Flujo de Trabajo de un Pedido", level=1)
    doc.add_paragraph(
        "Cada pedido sigue un flujo secuencial de etapas. Las dos primeras (Solicitud comercial "
        "y Validación administrativa) son obligatorias y bloqueadas. Las siguientes pueden "
        "personalizarse con plantillas o construirse manualmente."
    )

    doc.add_heading("12.1 Secuencia estándar", level=2)
    add_styled_table(doc,
        ["#", "Etapa", "Rol", "Auto-acción", "Siguiente paso"],
        [
            ["1", "Solicitud comercial", "Comercial", "Se completa al crear", "Pasa a admin"],
            ["2", "Validación administrativa", "Admin", "Se asigna a admin automáticamente", "Tras aprobar, libera etapa 3"],
            ["3", "Diseño técnico", "Diseño", "Empleado inicia manualmente", "Tras completar, libera etapa 4"],
            ["4", "Carpintería", "Carpintero", "Empleado inicia manualmente", "Tras completar, libera etapa 5"],
            ["5", "Tapicería", "Tapicería", "Solo si needsCushions=true", "Tras completar, libera etapa 6"],
            ["6", "Control de calidad", "Calidad", "Empleado inicia manualmente", "Tras completar, libera etapa 7"],
            ["7", "Despacho", "Despacho", "Empleado inicia manualmente", "Pedido → completed"],
        ])

    doc.add_heading("12.2 Derivación automática del estado", level=2)
    add_styled_table(doc,
        ["Condición", "Estado resultante"],
        [
            ["Todas las etapas completadas", "completed"],
            ["Alguna etapa bloqueada", "blocked"],
            ["Alguna etapa no-admin activa", "in_progress"],
            ["Etapa admin (índice 1) completada", "approved"],
            ["Ninguna de las anteriores", "pending_approval"],
        ])

    doc.add_heading("12.3 Cálculo de prioridad", level=2)
    add_styled_table(doc,
        ["Días restantes", "Prioridad", "Color"],
        [
            ["≤ 4 días", "Alta", "#ff6b91 (rojo)"],
            ["5 a 10 días", "Media", "#ffb36c (naranja)"],
            ["≥ 11 días", "Baja", "#7bffb6 (verde)"],
        ])

    doc.add_heading("12.4 Fórmulas de precios", level=2)
    doc.add_paragraph("Precio final = Costo estimado × 1.45 (margen del 45%)")
    doc.add_paragraph("Costo de producción = Costo estimado (ingresado por comercial)")
    doc.add_paragraph("Margen esperado = Precio final − Costo de producción")
    section_break(doc)


def ch13_roles(doc):
    doc.add_heading("13. Sistema de Roles y Permisos", level=1)
    doc.add_paragraph(
        "La aplicación implementa 7 roles con acceso diferenciado a vistas y acciones."
    )
    add_styled_table(doc,
        ["Rol", "Cuenta demo", "Color", "Responsabilidades principales"],
        [
            ["Comercial", "Laura Díaz", "Ámbar", "Captura de solicitudes, relación con cliente"],
            ["Administración", "Mateo Ruiz", "Teal", "Aprobación, precios, asignaciones, pagos"],
            ["Diseño", "Camila Soto", "Azul", "Planos y diseños técnicos"],
            ["Carpintería", "Javier Peña", "Marrón", "Estructura y ensamblaje de madera"],
            ["Tapicería", "Daniela Ríos", "Púrpura", "Tapizado y acabados textiles"],
            ["Calidad", "Sofía Lara", "Slate", "Control de calidad pre-entrega"],
            ["Despacho", "Esteban Cruz", "Dark Slate", "Logística y entrega al cliente"],
        ])

    doc.add_heading("Matriz de permisos", level=2)
    add_styled_table(doc,
        ["Acción", "Comercial", "Admin", "Diseño", "Carp.", "Tapic.", "Calidad", "Desp."],
        [
            ["Crear solicitud", "✅", "✅", "—", "—", "—", "—", "—"],
            ["Aprobar pedido", "—", "✅", "—", "—", "—", "—", "—"],
            ["Editar precios/margen", "—", "✅", "—", "—", "—", "—", "—"],
            ["Reasignar etapas", "—", "✅", "—", "—", "—", "—", "—"],
            ["Marcar pagos", "—", "✅", "—", "—", "—", "—", "—"],
            ["Cambiar prioridad", "—", "✅", "—", "—", "—", "—", "—"],
            ["Reordenar etapas", "—", "✅", "—", "—", "—", "—", "—"],
            ["Iniciar su etapa", "—", "—", "✅", "✅", "✅", "✅", "✅"],
            ["Completar su etapa", "—", "—", "✅", "✅", "✅", "✅", "✅"],
            ["Ver pagos propios", "✅", "✅", "✅", "✅", "✅", "✅", "✅"],
            ["Generar PDF", "✅", "✅", "✅", "✅", "✅", "✅", "✅"],
            ["Ver notificaciones", "✅", "✅", "✅", "✅", "✅", "✅", "✅"],
        ])
    section_break(doc)


def ch14_notificaciones(doc):
    doc.add_heading("14. Sistema de Notificaciones", level=1)

    doc.add_heading("14.1 Disparadores automáticos", level=2)
    add_styled_table(doc,
        ["Evento", "Destinatarios", "Prioridad", "Canales"],
        [
            ["Pedido creado", "Administrador", "Media", "In-app, Push, Email"],
            ["Pedido aprobado", "Equipo completo", "Alta", "In-app, Push, Email"],
            ["Etapa iniciada", "Admin + responsable", "Media", "In-app, Push, Email"],
            ["Etapa completada", "Admin + siguiente responsable", "Media", "In-app, Push, Email"],
            ["Pedido completado", "Equipo completo", "Alta", "In-app, Push, Email"],
            ["PDF generado", "Creador", "Baja", "In-app"],
        ])

    doc.add_heading("14.2 Plataformas soportadas", level=2)
    add_styled_table(doc,
        ["Plataforma", "API utilizada", "Permisos"],
        [
            ["iOS", "Expo Notifications (scheduleNotificationAsync)", "Solicita permiso explícito"],
            ["Android", "Expo Notifications (scheduleNotificationAsync)", "Automático en SDK 33+"],
            ["Web", "window.Notification (Web Notifications API)", "Solicita Notification.requestPermission()"],
        ])
    section_break(doc)


def ch15_pdf(doc):
    doc.add_heading("15. Generación de PDF y Correos", level=1)

    doc.add_heading("15.1 Estructura del PDF", level=2)
    doc.add_paragraph(
        "El PDF se genera desde una plantilla HTML profesional con colores neutros cálidos "
        "(#f6f1e7, #fff9ef) y tipografía limpia."
    )
    add_styled_table(doc,
        ["Sección", "Contenido"],
        [
            ["Hero", "Referencia, nombre del mueble, cliente, ciudad, fecha entrega"],
            ["Especificaciones", "Material, tamaño, almohadas, notas"],
            ["Costos", "Producción, precio final, margen esperado"],
            ["Flujo de trabajo", "Tabla: etapa, rol, responsable, estado, pago"],
        ])

    doc.add_heading("15.2 Envío de correo", level=2)
    add_styled_table(doc,
        ["Plataforma", "Método", "Adjunto PDF"],
        [
            ["iOS/Android", "MailComposer.composeAsync() — UI nativa", "Sí (archivo temporal)"],
            ["Web", "mailto: con asunto y cuerpo pre-formateado", "No (limitación del protocolo)"],
            ["Backend", "Cloud Function + Nodemailer (automático)", "Preparado (no implementado)"],
        ])
    section_break(doc)


def ch16_persistencia(doc):
    doc.add_heading("16. Persistencia y Sincronización", level=1)

    doc.add_heading("16.1 Estrategia Local-First", level=2)
    add_styled_table(doc,
        ["Operación", "Destino primario", "Destino secundario", "Fallback"],
        [
            ["Lectura", "Estado en memoria", "—", "—"],
            ["Escritura", "AsyncStorage (debounce 300ms)", "Firestore (background)", "Solo local si Firebase no disponible"],
            ["Carga inicial", "AsyncStorage", "Firestore", "Datos seed"],
        ])

    doc.add_heading("16.2 Cola de sincronización", level=2)
    doc.add_paragraph(
        "Cada mutación crea un SyncJob con estado 'pending'. Al sincronizar exitosamente "
        "con Firestore, el estado cambia a 'synced'. Si falla, queda 'failed' y se puede "
        "reintentar. El panel Centro Avanzado muestra la cola completa."
    )

    doc.add_heading("16.3 Salud del sistema", level=2)
    add_styled_table(doc,
        ["Indicador", "Healthy", "Warning", "Critical"],
        [
            ["Cola de sincronización", "≤ 3 items", "4 a 8 items", "> 8 items"],
            ["Correos pendientes", "≤ 3", "> 3", "—"],
            ["Pedidos vencidos", "0", "1 a 2", "> 2"],
        ])
    section_break(doc)


def ch17_scripts(doc):
    doc.add_heading("17. Scripts y Herramientas de Desarrollo", level=1)
    add_styled_table(doc,
        ["Script", "Comando", "Descripción"],
        [
            ["Desarrollo", "npm start", "Servidor Expo con menú interactivo"],
            ["Web", "npm run web", "Inicia app en navegador"],
            ["Android", "npm run android", "Inicia en emulador/dispositivo Android"],
            ["iOS", "npm run ios", "Inicia en simulador iOS"],
            ["Type check", "npm run typecheck", "Verifica tipos sin compilar (tsc --noEmit)"],
            ["Export web", "npm run export:web", "Build estático para deploy web"],
            ["Docs Word", "npm run docs:word", "Genera este documento Word"],
            ["Serve web", "pwsh scripts/serve-web.ps1", "Sirve build web exportado"],
        ])
    section_break(doc)


def ch18_mejoras(doc):
    doc.add_heading("18. Catálogo de 250 Mejoras", level=1)
    doc.add_paragraph(
        "La aplicación incluye un catálogo integrado de 250 mejoras organizadas en 10 categorías "
        "de 25 ítems cada una. Se visualizan en la pestaña Mejoras con estado (Activa/Preparada)."
    )

    path = ROOT / "src" / "constants" / "improvements.ts"
    if path.exists():
        content = path.read_text(encoding="utf-8")
        blocks = content.split("items: [")
        category_titles = re.findall(r"title: '([^']+)'", content)
        category_states = re.findall(r"state: '([^']+)'", content)

        for index, title in enumerate(category_titles):
            state = category_states[index] if index < len(category_states) else "N/D"
            doc.add_heading(f"{index + 1}. {title} ({state})", level=2)
            if index < len(blocks) - 1:
                items_block = blocks[index + 1].split("],", 1)[0]
                items = re.findall(r"'([^']+)'", items_block)
                for item in items:
                    doc.add_paragraph(item, style="List Bullet")
    section_break(doc)


def ch19_despliegue(doc):
    doc.add_heading("19. Guía de Despliegue", level=1)

    doc.add_heading("19.1 Despliegue Web", level=2)
    for step in [
        "1. Ejecutar: npm run export:web",
        "2. Los archivos estáticos se generan en la carpeta dist/",
        "3. Subir el contenido de dist/ a cualquier hosting estático (Vercel, Netlify, S3, etc.)",
        "4. O servir localmente: pwsh scripts/serve-web.ps1",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("19.2 Despliegue Móvil (EAS Build)", level=2)
    for step in [
        "1. Instalar EAS CLI: npm install -g eas-cli",
        "2. Configurar proyecto: eas build:configure",
        "3. Build Android: eas build --platform android",
        "4. Build iOS: eas build --platform ios",
        "5. Submit a tiendas: eas submit --platform [android|ios]",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("19.3 Despliegue Firebase", level=2)
    for step in [
        "1. Instalar Firebase CLI: npm install -g firebase-tools",
        "2. Login: firebase login",
        "3. Desplegar reglas e índices: firebase deploy --only firestore",
        "4. Desplegar Cloud Functions: firebase deploy --only functions",
        "5. Verificar en Firebase Console que las colecciones estén creadas",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("19.4 Variables de entorno requeridas", level=2)
    add_styled_table(doc,
        ["Variable", "Descripción", "Ejemplo"],
        [
            ["EXPO_PUBLIC_FIREBASE_API_KEY", "API Key del proyecto Firebase", "AIzaSyXXXXXX"],
            ["EXPO_PUBLIC_FIREBASE_AUTH_DOMAIN", "Dominio de autenticación", "proyecto.firebaseapp.com"],
            ["EXPO_PUBLIC_FIREBASE_PROJECT_ID", "ID del proyecto", "mi-proyecto-12345"],
            ["EXPO_PUBLIC_FIREBASE_STORAGE_BUCKET", "Bucket de Storage", "proyecto.appspot.com"],
            ["EXPO_PUBLIC_FIREBASE_MESSAGING_SENDER_ID", "ID del sender de FCM", "1234567890"],
            ["EXPO_PUBLIC_FIREBASE_APP_ID", "ID de la aplicación web", "1:123:web:abc"],
        ])
    section_break(doc)


def ch20_glosario(doc):
    doc.add_heading("20. Glosario Técnico", level=1)
    add_styled_table(doc,
        ["Término", "Definición"],
        [
            ["Expo", "Framework para React Native que simplifica el desarrollo multiplataforma"],
            ["React Native", "Framework JavaScript para crear apps nativas usando React"],
            ["Firestore", "Base de datos NoSQL en tiempo real de Firebase"],
            ["Cloud Functions", "Funciones serverless que se ejecutan en respuesta a eventos"],
            ["AsyncStorage", "API de almacenamiento local key-value para React Native"],
            ["Glassmorphism", "Estilo UI con paneles semi-transparentes y desenfoque"],
            ["Debounce", "Técnica que retrasa la ejecución hasta que cesen los eventos"],
            ["Batch write", "Escritura atómica de múltiples documentos en Firestore"],
            ["SMTP", "Protocolo estándar para envío de correo electrónico"],
            ["Nodemailer", "Biblioteca Node.js para envío de emails vía SMTP"],
            ["TypeScript strict", "Modo de compilación con máxima verificación de tipos"],
            ["CRUD", "Create, Read, Update, Delete — operaciones básicas de datos"],
            ["KPI", "Key Performance Indicator — indicador clave de rendimiento"],
            ["COP", "Peso colombiano — moneda utilizada en la aplicación"],
            ["EAS", "Expo Application Services — plataforma de build y deploy"],
            ["OAuth", "Protocolo de autorización para acceso a APIs de terceros"],
            ["Webhook", "Llamada HTTP desencadenada por un evento (similar a Cloud Functions)"],
            ["SLA", "Service Level Agreement — acuerdo de nivel de servicio"],
            ["CRDT", "Conflict-free Replicated Data Type — estructura para sync multi-usuario"],
        ])


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> None:
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    cover_page(doc)
    table_of_contents(doc)
    ch01_resumen(doc)
    ch02_arquitectura(doc)
    ch03_estructura(doc)
    ch04_stack(doc)
    ch05_modelo_datos(doc)
    ch06_app_tsx(doc)
    ch07_servicios(doc)
    ch08_utilidades(doc)
    ch09_datos_config(doc)
    ch10_firebase(doc)
    ch11_tema(doc)
    ch12_flujo(doc)
    ch13_roles(doc)
    ch14_notificaciones(doc)
    ch15_pdf(doc)
    ch16_persistencia(doc)
    ch17_scripts(doc)
    ch18_mejoras(doc)
    ch19_despliegue(doc)
    ch20_glosario(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUTPUT))
    print(f"Documento generado: {OUTPUT}")


if __name__ == "__main__":
    main()
